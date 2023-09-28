from os import getcwd
from statistics import median
from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import argparse
import aspose.words as aw


pars = argparse.ArgumentParser()

pars.add_argument("path", type=str, help="Path to the table")
pars.add_argument("--name", "-n", help="Bird/region name, must be in '' ", dest="name", default=None)
pars.add_argument("-s", action="store_true", dest="stat", help="get statistic")
pars.add_argument("-g", action="store_true", dest="plot", help="get plot")
pars.add_argument("-c", "--compare", dest="compare", action="store", help="Compare first arg with smth", default=None)
pars.add_argument("-l", action="store_true", dest="list", help="show list of birds/regions")

args = pars.parse_args()

bird_list = []
teritory_list = []
available_plots = []

done_table = []

cur_dir = getcwd()
file_to_write = Document()
margin = 1
sections = file_to_write.sections

for section in sections:
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def prepare(path: str) -> None:
    global bird_list
    global teritory_list
    global done_table

    table = pd.read_csv(path)

    year_list = table['survey_year'].unique()
    bird_list = table['bird_type'].unique()
    teritory_list = [val for val in table['bioregions'].unique() if type(val) is not float]

    file_to_write.add_paragraph(f"\nВсего строк в таблице - {len(table)}\n"
                                f"Количество видов птиц - {len(table['bird_type'].unique())}\n"
                                f"Данные собраны с {int(min(year_list))} по {int(max(year_list))} год\n")

    short_table = [table.query("urban_rural == 'Urban' & bird_count != 0"),
                   table.query("urban_rural == 'Rural' & bird_count != 0")]

    done_table = [short_table[0].groupby(['bioregions', 'bird_type'])['bird_count'].sum(),
                  short_table[1].groupby(['bioregions', 'bird_type'])['bird_count'].sum()]


def analise_prepare(name: str) -> tuple:
    rur_amount = 0
    urb_amount = 0
    cur_bird_dict = {}

    for region in teritory_list:
        try:
            urb_amount += done_table[0][region][name]
            if region not in cur_bird_dict:
                cur_bird_dict.update({region: done_table[0][region][name]})
            else:
                cur_bird_dict[region] += done_table[0][region][name]
        except KeyError:
            pass
        try:
            rur_amount += done_table[1][region][name]
            if region not in cur_bird_dict:
                cur_bird_dict.update({region: done_table[1][region][name]})
            else:
                cur_bird_dict[region] += done_table[1][region][name]
        except KeyError:
            pass

    return cur_bird_dict, rur_amount, urb_amount


def analise_region(name: str) -> tuple:
    rur_amount = 0
    urb_amount = 0
    different_birds = {}

    for cur_bird in bird_list:
        try:
            if cur_bird not in different_birds:
                different_birds.update({cur_bird: done_table[0][name][cur_bird]})
            else:
                different_birds[cur_bird] += done_table[0][name][cur_bird]
        except KeyError:
            pass

        try:
            if cur_bird not in different_birds:
                different_birds.update({cur_bird: done_table[1][name][cur_bird]})
            else:
                different_birds[cur_bird] += done_table[1][name][cur_bird]
        except KeyError:
            pass

        try:
            rur_amount += done_table[1][name][cur_bird]
        except KeyError:
            pass

        try:
            urb_amount += done_table[0][name][cur_bird]
        except KeyError:
            pass

    if rur_amount > urb_amount:
        pop_territory = "Rural"
    elif urb_amount > rur_amount:
        pop_territory = "Urban"
    else:
        pop_territory = "Urban and Rural"

    return (len(different_birds.keys()), sum(different_birds.values()), [key for key, val in different_birds.items()
                                                                         if val == max(different_birds.values())][0],
            max(different_birds.values()), median(different_birds.values()), pop_territory, different_birds, rur_amount,
            urb_amount)


def analise(name: str) -> None:
    if name in bird_list:
        response = analise_prepare(name)
        cur_bird_dict = response[0]
        rur_amount = response[1]
        urb_amount = response[2]
        cur_bird_list = cur_bird_dict.values()

        if rur_amount > urb_amount:
            bird_fav_ter = "Rural"
        elif urb_amount > rur_amount:
            bird_fav_ter = "Urban"
        else:
            bird_fav_ter = "Rural and Urban"

        if not args.compare:
            file_to_write.add_paragraph(f"\nВсего особоей {name} на исследованных территориях, "
                                        f"согласно таблице насчитывается - {sum(cur_bird_list)}\n"
                                        f"Наибольшее количество было зафиксировано в "
                                        f"{[key for key, val in cur_bird_dict.items() if val == max(cur_bird_list)][0]}"
                                        f", "
                                        f"там обитает {max(cur_bird_list)} {name}\n"
                                        f"В среднем в регионах с {name} встречается {median(cur_bird_list)} {name}\n"
                                        f"Предпочтительная среда обитания - {bird_fav_ter}")
        else:
            file_to_write.add_paragraph("Смотрите статистику в разделе сравнения")

    elif name in teritory_list:
        args_tuple = analise_region(name)
        if not args.compare:
            file_to_write.add_paragraph(f"\nСогласно таблице, всего в {name} обитает {args_tuple[0]} "
                                        f"различных видов птиц\n"
                                        f"Всего на территории {name} проживает {args_tuple[1]} "
                                        f"особей, принадлежащих этим видам\n"
                                        f"Чаще всего встречается "
                                        f"{args_tuple[2]}\n"
                                        f"Их там {args_tuple[3]}\n"
                                        f"В среднем можно встреить "
                                        f"{args_tuple[4]} особей одного вида\n"
                                        f"Птицы в данном регионе предпочитают {args_tuple[5]} территории\n")
        else:
            file_to_write.add_paragraph("Смотрите статистику в разделе сравнения")


def plot_prepare(name: str) -> tuple:
    bird_in_region = {}
    bird_territory = {}
    for region in teritory_list:
        if region not in bird_in_region:
            bird_in_region.update({region: 0})
        try:
            bird_in_region[region] += done_table[0][region][name]
        except KeyError:
            pass
        try:
            bird_in_region[region] += done_table[1][region][name]
        except KeyError:
            pass

        if region not in bird_territory:
            bird_territory.update({region: {"Urban": 0, "Rural": 0}})
        try:
            bird_territory[region]["Rural"] += done_table[1][region][name]
        except KeyError:
            pass
        try:
            bird_territory[region]["Urban"] += done_table[0][region][name]
        except KeyError:
            pass

    return bird_in_region, bird_territory


def plot_draw(region: list, bird_amount: list, territory_name: list,
              territory_name_short: list, rur_territory: list, urb_territory: list, name: str) -> None:
    available_plots.clear()
    width = 0.4
    w_list = np.arange(len(territory_name))

    plt.figure(figsize=(11, 9))
    plt.rcParams['font.size'] = '9'
    plt.subplot()
    plt.title(f"{name} in regions")
    plt.pie(bird_amount, autopct='%1.0f%%',
            startangle=90)
    plt.axis('equal')
    plt.legend(region, bbox_to_anchor=(0.8, 1), loc='upper left', borderaxespad=0)

    plt.savefig('plot_bird_1.png')
    available_plots.append('plot_bird_1.png')

    plt.figure(figsize=(11, 9))
    plt.rcParams['font.size'] = '9'
    plt.subplot()
    plt.title("Rural/Urban compare")
    plt.xticks(w_list, territory_name_short)
    plt.xlabel("Region name")
    plt.ylabel("Amount")
    plt.bar(w_list - (width / 2), rur_territory, label="Rural", width=width)
    plt.bar(w_list + (width / 2), urb_territory, label="Urban", width=width)
    plt.legend()

    plt.savefig('plot_bird_2.png')
    available_plots.append('plot_bird_2.png')


def plot_draw_region(birds_info: dict, name: str, rur_amount: int, urb_amount: int):
    available_plots.clear()
    sorted_birds_info = dict(sorted(birds_info.items(), key=lambda item: item[1], reverse=True))
    bird_name = [key for key in sorted_birds_info.keys()][:10]
    bird_name_short = ["".join([word[0] for word in bird.split()]) for bird in bird_name]
    most_pop_amount = [val for val in sorted_birds_info.values()][:10]
    bar_colors = ['red', 'yellow', 'green', 'gray', 'blue', 'black', 'orange', 'purple', 'pink', 'brown']

    plt.figure(figsize=(11, 9))
    plt.rcParams['font.size'] = '9'
    plt.subplot(1, 2, 1)
    plt.title(f"Top 10 birds in region {name}")
    plt.bar(bird_name_short, most_pop_amount, label=bird_name_short, color=bar_colors)
    plt.legend(bird_name)

    plt.subplot(1, 2, 2)
    plt.title("Rural/Urban compare")
    plt.xlabel("Territory")
    plt.ylabel("Amount")
    plt.bar(['Rural'], rur_amount, label="Rural", width=0.4)
    plt.bar(['Urban'], urb_amount, label="Urban", width=0.4)
    plt.legend()

    plt.savefig('plot_region.png')
    available_plots.append('plot_region.png')


def plot_show(name: str) -> None:
    if name in bird_list:
        response = plot_prepare(name)
        bird_in_region = response[0]
        bird_territory = response[1]

        region = [key for key in bird_in_region.keys() if bird_in_region[key] != 0]
        bird_amount = [val for val in bird_in_region.values() if val != 0]
        territory_name = [key for key in bird_territory.keys() if bird_territory[key]["Rural"] != 0 or
                          bird_territory[key]["Urban"] != 0]
        territory_name_short = [" ".join([word[0] for word in region.split()]) for region in territory_name]
        rur_territory = [bird_territory[region]["Rural"] for region in territory_name]
        urb_territory = [bird_territory[region]["Urban"] for region in territory_name]

        plot_draw(region, bird_amount, territory_name, territory_name_short, rur_territory, urb_territory, name)

    elif name in teritory_list:
        args_tuple = analise_region(name)
        plot_draw_region(args_tuple[6], name, args_tuple[7], args_tuple[8])


def compare(name_f: str, name_s: str) -> None:
    if name_f and name_s in bird_list:
        first_response = analise_prepare(name_f)
        second_response = analise_prepare(name_s)
        cur_bird_dict = [first_response[0], second_response[0]]
        rur_amount = [first_response[1], second_response[1]]
        urb_amount = [first_response[2], second_response[2]]
        cur_bird_list = [cur_bird_dict[0].values(), cur_bird_dict[1].values()]

        if rur_amount[0] > urb_amount[0]:
            bird_fav_ter_1 = "Rural"
        elif urb_amount[0] > rur_amount[0]:
            bird_fav_ter_1 = "Urban"
        else:
            bird_fav_ter_1 = "Rural and Urban"

        if rur_amount[1] > urb_amount[1]:
            bird_fav_ter_2 = "Rural"
        elif urb_amount[1] > rur_amount[1]:
            bird_fav_ter_2 = "Urban"
        else:
            bird_fav_ter_2 = "Rural and Urban"

        max_region_1 = [key for key, val in cur_bird_dict[0].items() if val == max(cur_bird_list[0])][0]
        max_region_2 = [key for key, val in cur_bird_dict[1].items() if val == max(cur_bird_list[1])][0]

        file_to_write.add_paragraph(f"\nСравнение {name_f} и {name_s}\n" 
                                    f"\nКоличество особей вида {name_f} согласно таблице - {sum(cur_bird_list[0])}\n" 
                                    f"Количество особей вида {name_s} согласно таблице - {sum(cur_bird_list[1])}\n" 
                                    f"Наибольшее количество {name_f} было зафиксировано в " 
                                    f"{max_region_1}, " 
                                    f"там обитает {max(cur_bird_list[0])} {name_f}\n" 
                                    f"Наибольшее количество {name_s} было зафиксировано в " 
                                    f"{max_region_2}, " 
                                    f"там обитает {max(cur_bird_list[1])} {name_s}\n" 
                                    f"В среднем в регионах с {name_f} встречается {median(cur_bird_list[0])} {name_f}\n" 
                                    f"В среднем в регионах с {name_s} встречается {median(cur_bird_list[1])} {name_s}\n" 
                                    f"Предпочтительная среда обитания для {name_f} - {bird_fav_ter_1}\n" 
                                    f"Предпочтительная среда обитания для {name_s} - {bird_fav_ter_2}\n")
        if sum(cur_bird_list[0]) > sum(cur_bird_list[1]):
            file_to_write.add_paragraph(f"Было обнаружено больше {name_f} чем {name_s} на "
                                        f"{abs(max(cur_bird_list[0]) - max(cur_bird_list[1]))} особей")
        else:
            file_to_write.add_paragraph(f"Было обнаружено больше {name_s} чем {name_f} на "
                                        f"{abs(max(cur_bird_list[0]) - max(cur_bird_list[1]))} особей")
        if median(cur_bird_list[0]) > median(cur_bird_list[1]):
            file_to_write.add_paragraph(f"В среднем {name_f} больше чем {name_s}")
        else:
            file_to_write.add_paragraph(f"В среднем {name_s} больше чем {name_f}")
        if max_region_1 == max_region_2:
            file_to_write.add_paragraph(f"{name_f} и {name_s} имеют наибольшую популяцию в одном регионе - "
                                        f"{max_region_1}")
        else:
            file_to_write.add_paragraph(f"{name_f} и {name_s} имеют наибольшую популяцию в разных регионах,"
                                        f"{name_f} в {max_region_1}, {name_s} в {max_region_2}")
        if bird_fav_ter_1 == bird_fav_ter_2:
            file_to_write.add_paragraph(f"{name_f} и {name_s} одинаково предпочитают {bird_fav_ter_1}")
        else:
            file_to_write.add_paragraph(f"{name_f} предпочитает территорию - {bird_fav_ter_1}, a "
                                        f"{name_s} предпочитает территорию - {bird_fav_ter_2}")
        if args.plot:
            file_to_write.add_heading(f'Графики для {name_s}')
            plot_show(name_s)
            for pict in available_plots:
                file_to_write.add_picture(pict, width=Inches(8))

    elif name_f and name_s in teritory_list:
        args_tuple_f = analise_region(name_f)
        args_tuple_s = analise_region(name_s)
        file_to_write.add_paragraph(f"\nСогласно таблице, всего в {name_f} обитает {args_tuple_f[0]} "
                                    f"различных видов птиц\n"
                                    f"Согласно таблице, всего в {name_s} обитает {args_tuple_s[0]} "
                                    f"различных видов птиц\n"
                                    f"Всего на территории {name_f} проживает {args_tuple_f[1]} "
                                    f"особей, принадлежащих этим видам\n"
                                    f"Чаще всего встречается "
                                    f"{args_tuple_f[2]}\n"
                                    f"Их там {args_tuple_s[3]}\n"
                                    f"Всего на территории {name_s} проживает {args_tuple_s[1]} "
                                    f"особей, принадлежащих этим видам\n"
                                    f"Чаще всего встречается "
                                    f"{args_tuple_s[2]}\n"
                                    f"Их там {args_tuple_s[3]}\n"
                                    f"В среднем можно встреить "
                                    f"{args_tuple_f[4]} особей одного вида в {name_f}\n"
                                    f"В среднем можно встреить "
                                    f"{args_tuple_s[4]} особей одного вида в {name_s}\n"
                                    f"Птицы в {name_f} предпочитают {args_tuple_f[5]} территории\n"
                                    f"Птицы в {name_s} предпочитают {args_tuple_s[5]} территории\n")
        if args_tuple_f > args_tuple_s:
            file_to_write.add_paragraph(f"Видов в регионе {name_f} больше, чем в {name_s}")
        else:
            file_to_write.add_paragraph(f"Видов в регионе {name_s} больше, чем в {name_f}")
        if args_tuple_f[1] > args_tuple_s[1]:
            file_to_write.add_paragraph(f"Общее количество особей в регионе {name_f} больше чем в {name_s}")
        else:
            file_to_write.add_paragraph(f"Общее количество особей в регионе {name_s} больше чем в {name_f}")
        if args_tuple_f[2] == args_tuple_s[2]:
            file_to_write.add_paragraph(f"Наиболее часто встречается в обоих регионах встречается "
                                        f"один и тот же вид - {args_tuple_f[2]}")
        else:
            file_to_write.add_paragraph(f"Наиболее часто встречающиеся виды различается в {name_f} и {name_s}"
                                        f",\n в {name_f} - {args_tuple_f[2]}, в {name_s} - {args_tuple_s[2]}")
        if args_tuple_f[4] > args_tuple_s[4]:
            file_to_write.add_paragraph(f"В среднем птиц в регоине {name_f} больше, чем в {name_s}")
        else:
            file_to_write.add_paragraph(f"В среднем птиц в регоине {name_s} больше, чем в {name_f}")
        if args_tuple_f[5] == args_tuple_s[6]:
            file_to_write.add_paragraph(f"Птицы в {name_f} и {name_s} предпочитают "
                                        f"одинаковые территории - {args_tuple_f[5]}")
        else:
            file_to_write.add_paragraph(f"Птицы в {name_f} и {name_s} предпочитают "
                                        f"разные территории,\n в {name_f} - {args_tuple_f[5]}, "
                                        f"а в {name_s} - {args_tuple_s[5]}")
        if args.plot:
            file_to_write.add_heading(f'Графики для {name_s}')
            plot_show(name_s)
            for pict in available_plots:
                file_to_write.add_picture(pict, width=Inches(8))
    else:
        print("Wrong request")


def main():
    if args.stat or args.plot or args.compare:
        file_to_write.add_heading('Общая информация')
        prepare(args.path)

    if args.stat:
        try:
            if args.name is not None:
                file_to_write.add_heading('Статистика')
                analise(args.name)
                print("Analise done!")
        except ZeroDivisionError:
            print("Wrong name")

    if args.compare is not None:
        file_to_write.add_heading('Сравнение')
        compare(args.name, args.compare)
        print("Compare done!")

    if args.list:
        bird_table = pd.read_csv(args.path)
        bird_name_list = list(bird_table['bird_type'].unique())
        bird_name_list.sort()
        print(*bird_name_list, sep=";\n")
        print("\n")
        print(*[val for val in bird_table['bioregions'].unique() if type(val) is not float], sep=";\n")

    if args.plot:
        if args.name is not None:
            file_to_write.add_heading(f'Графики для {args.name}')
            plot_show(args.name)
            print("Plots done!")
        else:
            print("need bird/region name")

    if len(available_plots) != 0:
        for pic in available_plots:
            file_to_write.add_picture(pic, width=Inches(8))
    file_to_write.save(cur_dir + "/report.docx")
    doc = aw.Document(cur_dir + "/report.docx")
    doc.save(cur_dir + "/report.pdf")
    print("Все готово!",
          "\nПосмотреть отчет можно в формате docx и pdf в текущей рабочей директории, файл report.docx/report.pdf")


if __name__ == "__main__":
    main()
